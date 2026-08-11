import base64
from pathlib import Path

import docx
from openai import APIError, APITimeoutError, RateLimitError

from src.constants.models import GPT4O
from src.infrastructure.parsing.openai_client import get_openai_client

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
DOCX_EXTENSIONS = {".docx"}
VALID_EXTENSIONS = IMAGE_EXTENSIONS | DOCX_EXTENSIONS
MAX_OUTPUT_TOKENS = 4096

SYSTEM_PROMPT = (
    "Sos un asistente experto en digitalizar documentos legales. Tu unica tarea es "
    "transcribir fielmente el texto completo de la imagen de un contrato, preservando "
    "la jerarquia original: numeracion de clausulas, secciones, subsecciones, titulos "
    "y parrafos. No resumas, no interpretes, no agregues comentarios ni corrijas "
    "errores del documento original. Si una parte es ilegible, marcala como [ILEGIBLE] "
    "en vez de inventar contenido."
)


class DocumentValidationError(ValueError):
    pass


class VisionParsingError(RuntimeError):
    pass


class DocxParsingError(RuntimeError):
    pass


def _validate_document_path(document_path: str) -> Path:
    path = Path(document_path)
    if not path.exists():
        raise DocumentValidationError(f"El archivo no existe: {document_path}")
    if not path.is_file():
        raise DocumentValidationError(f"La ruta no es un archivo: {document_path}")
    if path.suffix.lower() not in VALID_EXTENSIONS:
        valid = ", ".join(sorted(VALID_EXTENSIONS))
        raise DocumentValidationError(
            f"Formato no soportado ({path.suffix}). Formatos validos: {valid}"
        )
    return path


def _encode_image_base64(path: Path) -> str:
    with path.open("rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")


def _mime_type(path: Path) -> str:
    return "image/png" if path.suffix.lower() == ".png" else "image/jpeg"


def _parse_image(path: Path) -> str:
    encoded_image = _encode_image_base64(path)
    mime = _mime_type(path)

    client = get_openai_client()
    try:
        response = client.chat.completions.create(
            model=GPT4O,
            temperature=0,
            max_tokens=MAX_OUTPUT_TOKENS,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Transcribi el texto completo de este contrato.",
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime};base64,{encoded_image}"},
                        },
                    ],
                },
            ],
        )
    except RateLimitError as error:
        raise VisionParsingError(
            f"Limite de rate excedido al parsear {path.name}: {error}"
        ) from error
    except APITimeoutError as error:
        raise VisionParsingError(f"Timeout al parsear {path.name}: {error}") from error
    except APIError as error:
        raise VisionParsingError(
            f"Error de la API de OpenAI al parsear {path.name}: {error}"
        ) from error

    text = response.choices[0].message.content
    if not text or not text.strip():
        raise VisionParsingError(f"El modelo no devolvio texto para {path.name}")
    return text.strip()


def _parse_docx(path: Path) -> str:
    try:
        document = docx.Document(str(path))
    except Exception as error:
        raise DocxParsingError(f"No se pudo abrir el archivo Word {path.name}: {error}") from error

    # El texto de las tablas se agrega al final porque python-docx expone
    # parrafos y tablas como colecciones separadas, no en orden de documento.
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    text = "\n".join(paragraphs + table_cells)

    if not text.strip():
        raise DocxParsingError(f"El documento Word no tiene texto: {path.name}")
    return text.strip()


def parse_contract_document(document_path: str) -> str:
    path = _validate_document_path(document_path)
    if path.suffix.lower() in DOCX_EXTENSIONS:
        return _parse_docx(path)
    return _parse_image(path)
