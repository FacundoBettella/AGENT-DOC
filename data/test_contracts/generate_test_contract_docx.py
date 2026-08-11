"""Genera un par de contratos de prueba en formato Word (requiere python-docx).

Mismo contenido que par1_servicios_simple (imagenes), pero como .docx, para
ejercitar el path de parsing de Word (extraccion directa de texto, sin Vision).

Uso: python data/test_contracts/generate_test_contract_docx.py
"""

from pathlib import Path

from docx import Document

OUTPUT_DIR = Path(__file__).parent / "par3_word_docx"
OUTPUT_DIR.mkdir(exist_ok=True)


def render_contract(title: str, intro: str | None, clauses: list[tuple[str, str]], path: Path) -> None:
    document = Document()
    document.add_heading(title, level=1)

    if intro:
        document.add_paragraph(intro)

    for header, body in clauses:
        document.add_heading(header, level=2)
        document.add_paragraph(body)

    document.save(str(path))


render_contract(
    title="CONTRATO DE PRESTACION DE SERVICIOS",
    intro=None,
    clauses=[
        (
            "PRIMERA. PARTES.",
            "El presente contrato de prestacion de servicios se celebra entre "
            "LegalMove S.A., en adelante EL CLIENTE, y Consultora Andina SRL, en "
            "adelante EL PRESTADOR.",
        ),
        (
            "SEGUNDA. OBJETO.",
            "EL PRESTADOR se compromete a brindar servicios de consultoria en "
            "materia de cumplimiento normativo a EL CLIENTE, de acuerdo a los "
            "terminos aqui establecidos.",
        ),
        (
            "TERCERA. MONTO.",
            "EL CLIENTE abonara a EL PRESTADOR la suma mensual de USD 1.500 (mil "
            "quinientos dolares estadounidenses), pagaderos dentro de los primeros "
            "cinco dias habiles de cada mes.",
        ),
        (
            "CUARTA. VIGENCIA.",
            "El presente contrato entrara en vigencia a partir de su firma y "
            "mantendra su validez hasta el dia 30 de junio de 2026, fecha de "
            "vencimiento del presente acuerdo.",
        ),
        (
            "QUINTA. CONFIDENCIALIDAD.",
            "Ambas partes se comprometen a mantener confidencialidad respecto de "
            "toda informacion intercambiada en el marco de este contrato.",
        ),
        (
            "SEXTA. JURISDICCION.",
            "Para cualquier controversia derivada del presente contrato, las "
            "partes se someten a la jurisdiccion de los tribunales ordinarios de "
            "la Ciudad Autonoma de Buenos Aires.",
        ),
    ],
    path=OUTPUT_DIR / "original.docx",
)

render_contract(
    title="ENMIENDA N.1 AL CONTRATO DE PRESTACION DE SERVICIOS",
    intro=(
        "Las partes acuerdan modificar el contrato de prestacion de servicios "
        "suscripto oportunamente, unicamente en los terminos que se detallan a "
        "continuacion. El resto de las clausulas del contrato original mantienen "
        "plena vigencia."
    ),
    clauses=[
        (
            "TERCERA. MONTO (MODIFICADA).",
            "EL CLIENTE abonara a EL PRESTADOR la suma mensual de USD 1.800 (mil "
            "ochocientos dolares estadounidenses), pagaderos dentro de los "
            "primeros cinco dias habiles de cada mes.",
        ),
        (
            "CUARTA. VIGENCIA (MODIFICADA).",
            "El presente contrato mantendra su validez hasta el dia 31 de "
            "diciembre de 2026, fecha de vencimiento del presente acuerdo.",
        ),
    ],
    path=OUTPUT_DIR / "amendment.docx",
)

print("Contratos de prueba (Word) generados en", OUTPUT_DIR)
